import io


def extract_text_from_brief(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if ext in ('docx', 'doc'):
        return _from_docx(file_bytes)
    elif ext == 'pdf':
        return _from_pdf(file_bytes)
    elif ext in ('xlsx', 'xls'):
        return _from_excel(file_bytes)
    elif ext in ('txt', 'text', ''):
        return file_bytes.decode('utf-8', errors='ignore')
    else:
        return file_bytes.decode('utf-8', errors='ignore')


def _from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return '\n'.join(parts)


def _from_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text.strip())
    return '\n'.join(parts)


def _from_excel(file_bytes: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f'[Sheet: {sheet.title}]')
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(' | '.join(cells))
    return '\n'.join(parts)
